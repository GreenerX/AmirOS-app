from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "marketing"
LOGO = ROOT / "ui" / "public" / "amiros-mark-v2-cropped.png"
DOCX_PATH = OUTPUT / "AmirOS_Product_and_Brand_Brief.docx"
PDF_PATH = OUTPUT / "AmirOS_One_Page_Product_Overview.pdf"

OCEAN = "175B7E"
OCEAN_DARK = "103D56"
MINT = "56DDB2"
MIST = "EEF6F9"
INK = "142B3A"
MUTED = "617484"
LINE = "D8E6ED"


def set_run_font(run, *, name="Aptos", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_paragraph(doc, text, *, size=11, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return paragraph


def style_existing_paragraph(paragraph, *, size=11, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(4)
    run = paragraph.add_run("AmirOS | Product and brand reference")
    set_run_font(run, size=8, color=MUTED)


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=15, color=OCEAN, bold=True)
    return paragraph


def marketing_card(table, row, title, body, column, fill="FFFFFF"):
    cell = table.cell(row, column)
    shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    set_run_font(run, size=10.5, color=OCEAN_DARK, bold=True)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    for run in paragraph.runs:
        set_run_font(run, size=9.2, color=MUTED)


def build_docx():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(0.78)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    add_footer(section)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    brand_line = document.add_paragraph()
    brand_line.paragraph_format.space_after = Pt(8)
    brand_line.add_run().add_picture(str(LOGO), width=Inches(0.52))
    brand_run = brand_line.add_run("  AmirOS")
    set_run_font(brand_run, size=15, color=OCEAN_DARK, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("Product + Brand Brief")
    set_run_font(title_run, size=29, color=INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run("A relationship-aware WhatsApp assistant for context, clarity, and better follow-through.")
    set_run_font(subtitle_run, size=13, color=MUTED)

    callout = document.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    cell = callout.cell(0, 0)
    shade(cell, MIST)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run("THE IDEA")
    set_run_font(run, size=8.5, color=OCEAN, bold=True)
    paragraph = cell.add_paragraph("Your relationships have context. Your WhatsApp should too.")
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, size=16, color=OCEAN_DARK, bold=True)

    add_heading(document, "What AmirOS does")
    add_paragraph(
        document,
        "AmirOS is a desktop command center linked to WhatsApp. It turns important conversation context into useful next steps while helping users reply in their own voice. Each chat can have its own relationship context, tone, instructions, and knowledge access.",
        size=10.7,
        color=INK,
        after=9,
    )

    benefits = document.add_table(rows=2, cols=2)
    set_table_geometry(benefits, [4680, 4680])
    marketing_card(benefits, 0, "Remember what matters", "Approved relationship knowledge and chat summaries keep key people, preferences, plans, and context easy to revisit.", 0, MIST)
    marketing_card(benefits, 0, "Reply like yourself", "Per-chat writing style, tone, custom instructions, and memory help suggestions feel personal instead of generic.", 1, "F8FBFC")
    marketing_card(benefits, 1, "Turn chat into action", "Plans, commitments, calendar suggestions, and replies that need attention are organized into a single reviewable queue.", 0, "F8FBFC")
    marketing_card(benefits, 1, "Stay in control", "Users choose what is saved, which knowledge a chat can use, and whether assistance is Off, Suggest, or Auto.", 1, MIST)

    add_heading(document, "Core product moments")
    moments = document.add_table(rows=6, cols=2)
    set_table_geometry(moments, [2450, 6910])
    for row, (label, description) in enumerate([
        ("Inbox", "Live WhatsApp conversations with media, voice notes, replies, reactions, and per-contact settings."),
        ("Intelligence", "A command center for knowledge, people, action queues, commitments, and searchable answer history."),
        ("Calendar", "Editable event suggestions discovered in chat, with review before confirmation and calendar export options."),
        ("Ask AmirOS", "A scoped, private assistant for questions about people, chats, commitments, and the calendar."),
        ("Personalization", "Writing-style learning, relationship options, custom guidance, and access controls for each chat."),
        ("Control center", "Assistant models, usage, WhatsApp linking, behavior rules, privacy controls, and themes in one place."),
    ]):
        label_cell, body_cell = moments.rows[row].cells
        shade(label_cell, MIST if row % 2 == 0 else "F8FBFC")
        shade(body_cell, "FFFFFF")
        label_p = label_cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        label_run = label_p.add_run(label)
        set_run_font(label_run, size=9.5, color=OCEAN_DARK, bold=True)
        body_p = body_cell.paragraphs[0]
        body_p.paragraph_format.space_after = Pt(0)
        body_run = body_p.add_run(description)
        set_run_font(body_run, size=9.2, color=INK)

    document.add_page_break()
    add_heading(document, "Who it is for")
    add_paragraph(document, "AmirOS is designed for people who run a meaningful part of their life through WhatsApp: founders, busy professionals, planners, and relationship-oriented people who want more context without maintaining a manual CRM.", size=10.7, color=INK, after=8)

    audience = document.add_table(rows=1, cols=3)
    set_table_geometry(audience, [3120, 3120, 3120])
    marketing_card(audience, 0, "The thoughtful operator", "Wants to keep commitments, follow through, and reply with full context.", 0, MIST)
    marketing_card(audience, 0, "The relationship builder", "Values personal conversations and wants help remembering details that make people feel seen.", 1, "F8FBFC")
    marketing_card(audience, 0, "The power WhatsApp user", "Needs a premium desktop layer for the chats that organize work, plans, family, and friends.", 2, MIST)

    add_heading(document, "Brand voice")
    voice_table = document.add_table(rows=4, cols=2)
    set_table_geometry(voice_table, [2450, 6910])
    for row, (label, description) in enumerate([
        ("Personality", "Premium, warm, intelligent, discreet, useful, personal, and confident."),
        ("Lead with", "The human benefit - remember more, reply better, keep life moving - then explain the intelligence."),
        ("Avoid", "Cold automation language, surveillance-like imagery, generic robots, neon cyberpunk, and exaggerated claims."),
        ("Make tangible", "Control: review, approve, configure, choose the scope, and personalize each conversation."),
    ]):
        left, right = voice_table.rows[row].cells
        shade(left, MIST if row % 2 == 0 else "F8FBFC")
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        set_run_font(lr, size=9.5, color=OCEAN_DARK, bold=True)
        rp = right.paragraphs[0]
        rp.paragraph_format.space_after = Pt(0)
        rr = rp.add_run(description)
        set_run_font(rr, size=9.2, color=INK)

    add_heading(document, "Messaging starter set")
    messages = document.add_table(rows=5, cols=2)
    set_table_geometry(messages, [2820, 6540])
    for row, (purpose, copy) in enumerate([
        ("Primary headline", "Your relationships have context. Your WhatsApp should too."),
        ("Alternative", "The relationship intelligence layer for WhatsApp."),
        ("Alternative", "Remember more. Reply better. Keep life moving."),
        ("Short product description", "AmirOS turns WhatsApp conversations into organized relationship context, better replies, and actionable plans - all from a premium desktop command center designed around your voice and your control."),
        ("Call to action", "Meet AmirOS | See your conversations differently | Turn chat into clarity | Join the early access list"),
    ]):
        left, right = messages.rows[row].cells
        shade(left, MIST if row % 2 == 0 else "F8FBFC")
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(purpose)
        set_run_font(lr, size=9.2, color=OCEAN_DARK, bold=True)
        rp = right.paragraphs[0]
        rp.paragraph_format.space_after = Pt(0)
        rr = rp.add_run(copy)
        set_run_font(rr, size=9.2, color=INK)

    add_heading(document, "Creative direction for Pomelli")
    add_paragraph(document, "Use a premium desktop-product aesthetic: deep ocean blue, quiet mint accents, soft white surfaces, crisp editorial typography, rounded cards, and generous space. Show meaningful interface details rather than abstract AI imagery. Pair product screens with a concise outcome-focused caption.", size=10.5, color=INK, after=6)
    creative_callout = document.add_table(rows=1, cols=1)
    set_table_geometry(creative_callout, [9360])
    cell = creative_callout.cell(0, 0)
    shade(cell, "F8FBFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SCREENSHOT SAFETY")
    set_run_font(r, size=8.4, color=OCEAN, bold=True)
    p = cell.add_paragraph("Use clean interface crops that do not show real contact names, phone numbers, private messages, or real event details. A safe starting point is the Assistant models, Assistant availability, and Color theme settings surface.")
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, size=9.2, color=MUTED)

    document.save(DOCX_PATH)


def pdf_paragraph(text, style):
    return Paragraph(text, style)


def build_pdf():
    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.48 * inch,
        bottomMargin=0.48 * inch,
        title="AmirOS - Product Overview",
        author="AmirOS",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=26, leading=29, textColor=colors.HexColor(f"#{INK}"), spaceAfter=4)
    subtitle_style = ParagraphStyle("Subtitle", parent=styles["Normal"], fontName="Helvetica", fontSize=11, leading=15, textColor=colors.HexColor(f"#{MUTED}"), spaceAfter=12)
    kicker_style = ParagraphStyle("Kicker", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8.5, leading=10, textColor=colors.HexColor(f"#{OCEAN}"), spaceAfter=5, uppercase=True)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontName="Helvetica", fontSize=9.2, leading=12.2, textColor=colors.HexColor(f"#{INK}"))
    feature_title = ParagraphStyle("FeatureTitle", parent=body_style, fontName="Helvetica-Bold", fontSize=9.6, leading=12, textColor=colors.HexColor(f"#{OCEAN_DARK}"), spaceAfter=3)
    feature_body = ParagraphStyle("FeatureBody", parent=body_style, textColor=colors.HexColor(f"#{MUTED}"), fontSize=8.6, leading=11.2)
    message_style = ParagraphStyle("Message", parent=body_style, fontName="Helvetica-Bold", fontSize=11.7, leading=15, textColor=colors.HexColor(f"#{OCEAN_DARK}"), alignment=TA_CENTER)
    caption_style = ParagraphStyle("Caption", parent=body_style, fontSize=8.2, leading=10, textColor=colors.HexColor(f"#{MUTED}"), alignment=TA_CENTER)

    story = []
    header = Table([[Image(str(LOGO), width=0.4 * inch, height=0.4 * inch), pdf_paragraph("<b>AmirOS</b>", ParagraphStyle("Brand", parent=body_style, fontName="Helvetica-Bold", fontSize=14, textColor=colors.HexColor(f"#{OCEAN_DARK}")))]], colWidths=[0.48 * inch, 6.8 * inch])
    header.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
    story += [header, Spacer(1, 12), pdf_paragraph("RELATIONSHIP INTELLIGENCE FOR WHATSAPP", kicker_style), pdf_paragraph("Your relationships have context. Your WhatsApp should too.", title_style), pdf_paragraph("AmirOS is a desktop command center that helps users remember important details, reply in their own voice, and turn conversations into useful next steps.", subtitle_style)]

    message = Table([[pdf_paragraph("<b>Remember more. Reply better. Keep life moving.</b>", message_style)]], colWidths=[7.35 * inch])
    message.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{MIST}")), ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{LINE}")), ("TOPPADDING", (0, 0), (-1, -1), 13), ("BOTTOMPADDING", (0, 0), (-1, -1), 13), ("LEFTPADDING", (0, 0), (-1, -1), 17), ("RIGHTPADDING", (0, 0), (-1, -1), 17)]))
    story += [message, Spacer(1, 13)]

    feature_data = [
        [pdf_paragraph("Relationship intelligence", feature_title), pdf_paragraph("Personalized replies", feature_title), pdf_paragraph("Calendar awareness", feature_title)],
        [pdf_paragraph("Build approved, useful context for people and group chats.", feature_body), pdf_paragraph("Match each chat’s writing style, tone, instructions, and memory.", feature_body), pdf_paragraph("Turn plans found in chat into editable, reviewable event suggestions.", feature_body)],
        [pdf_paragraph("Unified action queue", feature_title), pdf_paragraph("Ask AmirOS", feature_title), pdf_paragraph("User control", feature_title)],
        [pdf_paragraph("See replies, promises, plans, and knowledge that need a decision.", feature_body), pdf_paragraph("Ask scoped questions across selected people, chats, commitments, and calendar context.", feature_body), pdf_paragraph("Choose what is saved, which knowledge can be used, and whether help is Off, Suggest, or Auto.", feature_body)],
    ]
    features = Table(feature_data, colWidths=[2.45 * inch] * 3)
    features.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 1), colors.HexColor("#F8FBFC")),
        ("BACKGROUND", (0, 2), (-1, 3), colors.HexColor("#FFFFFF")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{LINE}")),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{LINE}")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]))
    story += [features, Spacer(1, 14)]

    story += [pdf_paragraph("BUILT FOR PEOPLE WHO RUN LIFE THROUGH WHATSAPP", kicker_style)]
    audience = Table([[pdf_paragraph("<b>Founders and busy professionals</b><br/>Keep plans, promises, and relationships from getting lost in chat.", body_style), pdf_paragraph("<b>Relationship-oriented planners</b><br/>Reply with more context and make the people in your life feel remembered.", body_style)]], colWidths=[3.675 * inch, 3.675 * inch])
    audience.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.white), ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{LINE}")), ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{LINE}")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 11), ("BOTTOMPADDING", (0, 0), (-1, -1), 11), ("LEFTPADDING", (0, 0), (-1, -1), 12), ("RIGHTPADDING", (0, 0), (-1, -1), 12)]))
    story += [audience, Spacer(1, 15)]

    story += [pdf_paragraph("BRAND FEEL", kicker_style), pdf_paragraph("Premium, warm, intelligent, discreet, and personal. AmirOS should feel like a thoughtful chief of staff for your relationships - never cold, creepy, or overly automated.", body_style), Spacer(1, 8), pdf_paragraph("Creative direction: ocean blue, quiet mint accents, soft white surfaces, crisp editorial typography, rounded cards, and clear product details. Avoid generic robots, neon cyberpunk, message clutter, and surveillance-like visuals.", caption_style)]

    document.build(story)


if __name__ == "__main__":
    OUTPUT.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
